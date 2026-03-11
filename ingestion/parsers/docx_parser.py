"""Parser for .docx files using python-docx.

Extracts:
  - Paragraphs (non-empty, stripped)
  - Tables (each row as a pipe-delimited string)

Each returned dict has the shape:
  {
      "text": str,
      "source": str,        # filename
      "section": str,       # nearest heading or "table"
      "page": None,         # .docx has no reliable page numbers
  }
"""

import logging
from pathlib import Path

import docx

logger = logging.getLogger(__name__)


def parse(file_path: Path) -> list[dict[str, str | None]]:
    """Extract text chunks from a .docx file.

    Args:
        file_path: Absolute path to the .docx file.

    Returns:
        List of chunk dicts ready for the chunker.
    """
    document = docx.Document(str(file_path))
    source = file_path.name
    chunks: list[dict[str, str | None]] = []
    current_heading = "Introduction"

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Track headings to use as section labels
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            current_heading = text
            continue

        chunks.append(
            {
                "text": text,
                "source": source,
                "section": current_heading,
                "page": None,
            }
        )

    # Extract tables — each row becomes a pipe-delimited string
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(
                    {
                        "text": row_text,
                        "source": source,
                        "section": "table",
                        "page": None,
                    }
                )

    logger.info("docx_parser: extracted %d chunks from '%s'", len(chunks), source)
    return chunks
